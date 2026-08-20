#!/usr/bin/env python3
"""Build the editable Word package for personal-data legal approval."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "docs" / "legal-approval-package"
OUTPUT = SOURCE_DIR / "steelprodukt-pd-legal-approval-package.docx"

SOURCE_FILES = [
    "README.md",
    "01-approval-order.md",
    "02-personal-data-processing-regulation.md",
    "03-access-control-regulation.md",
    "04-subject-and-authority-requests.md",
    "05-retention-legal-hold-destruction.md",
    "06-official-export-and-transfer.md",
    "07-incident-response-regulation.md",
    "08-backup-and-restore-regulation.md",
    "09-register-forms.md",
    "10-historical-consent-exception-act.md",
    "11-rkn-registry-verification-sheet.md",
    "12-public-documents-approval-sheet.md",
    "13-operational-forms.md",
    "14-public-site-legal-review.md",
    "15-confidentiality-undertaking.md",
    "16-public-legal-texts.md",
]

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x66, 0x66, 0x66)
ORANGE = RGBColor(0xEA, 0x5B, 0x0C)
BLACK = RGBColor(0x00, 0x00, 0x00)
LIGHT_ORANGE = "FCE9DF"


def set_cellless_rule(paragraph, color: str = "2E74B5", size: str = "12") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr_text, fld_end])


def set_font(run, name: str = "Calibri", size: float | None = None,
             color: RGBColor | None = None, bold: bool | None = None,
             italic: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ["Bullet Compact", "Number Compact", "Checklist"]:
        if name not in doc.styles:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = doc.styles[name]
        style.base_style = normal
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Document Label" not in doc.styles:
        label = doc.styles.add_style("Document Label", WD_STYLE_TYPE.PARAGRAPH)
    else:
        label = doc.styles["Document Label"]
    label.base_style = normal
    label.font.name = "Calibri"
    label.font.size = Pt(9)
    label.font.bold = True
    label.font.color.rgb = ORANGE
    label.paragraph_format.space_after = Pt(4)
    label.paragraph_format.keep_with_next = True


def add_numbering(doc: Document) -> tuple[int, int, int, int]:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(item.get(qn("w:abstractNumId"))) for item in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(item.get(qn("w:numId"))) for item in numbering.findall(qn("w:num"))]
    next_abstract = max(abstract_ids, default=0) + 1
    next_num = max(num_ids, default=0) + 1

    def create(abstract_id: int, num_id: int, bullet: bool) -> None:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "multilevel")
        abstract.append(multi)
        for level in range(3):
            lvl = OxmlElement("w:lvl")
            lvl.set(qn("w:ilvl"), str(level))
            start = OxmlElement("w:start")
            start.set(qn("w:val"), "1")
            lvl.append(start)
            num_fmt = OxmlElement("w:numFmt")
            num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
            lvl.append(num_fmt)
            lvl_text = OxmlElement("w:lvlText")
            if bullet:
                lvl_text.set(qn("w:val"), "•" if level == 0 else "–")
            else:
                lvl_text.set(qn("w:val"), ".".join(f"%{i + 1}" for i in range(level + 1)) + ".")
            lvl.append(lvl_text)
            lvl_jc = OxmlElement("w:lvlJc")
            lvl_jc.set(qn("w:val"), "left")
            lvl.append(lvl_jc)
            p_pr = OxmlElement("w:pPr")
            tabs = OxmlElement("w:tabs")
            tab = OxmlElement("w:tab")
            tab.set(qn("w:val"), "num")
            tab.set(qn("w:pos"), str(540 + level * 360))
            tabs.append(tab)
            p_pr.append(tabs)
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"), str(540 + level * 360))
            ind.set(qn("w:hanging"), "270")
            p_pr.append(ind)
            lvl.append(p_pr)
            abstract.append(lvl)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_id_el = OxmlElement("w:abstractNumId")
        abstract_id_el.set(qn("w:val"), str(abstract_id))
        num.append(abstract_id_el)
        numbering.append(num)

    create(next_abstract, next_num, True)
    create(next_abstract + 1, next_num + 1, False)
    return next_num, next_num + 1, next_abstract, next_abstract + 1


def new_num_instance(doc: Document, abstract_id: int) -> int:
    numbering = doc.part.numbering_part.element
    num_ids = [int(item.get(qn("w:numId"))) for item in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_id_el = OxmlElement("w:abstractNumId")
    abstract_id_el.set(qn("w:val"), str(abstract_id))
    num.append(abstract_id_el)
    for level in range(3):
        override = OxmlElement("w:lvlOverride")
        override.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:startOverride")
        start.set(qn("w:val"), "1")
        override.append(start)
        num.append(override)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id: int, level: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(min(level, 2)))
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])


INLINE = re.compile(r"(\*\*.+?\*\*|`.+?`)")


def add_inline(paragraph, text: str) -> None:
    position = 0
    for match in INLINE.finditer(text):
        if match.start() > position:
            set_font(paragraph.add_run(text[position:match.start()]))
        token = match.group(0)
        if token.startswith("**"):
            set_font(paragraph.add_run(token[2:-2]), bold=True)
        else:
            set_font(paragraph.add_run(token[1:-1]), name="Courier New", size=9.5,
                     color=DARK_BLUE)
        position = match.end()
    if position < len(text):
        set_font(paragraph.add_run(text[position:]))


def configure_page(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    set_font(p.add_run("ООО «ЭНЕРГОАЛЬЯНС»  |  ПАКЕТ ЛОКАЛЬНЫХ АКТОВ"),
             size=8.5, color=MUTED, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    set_font(p.add_run("ПРОЕКТ — НЕ УТВЕРЖДЕНО  |  стр. "), size=8.5,
             color=MUTED)
    add_page_field(p)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    set_font(p.add_run("ПАКЕТ ЛОКАЛЬНЫХ АКТОВ"), size=23, bold=True,
             color=BLACK)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    set_font(p.add_run("Обработка и защита персональных данных"), size=14,
             color=DARK_BLUE)

    metadata = [
        ("Оператор", "ООО «ЭНЕРГОАЛЬЯНС»"),
        ("ИНН / ОГРН", "6732110789 / 1156733014657"),
        ("Сайт", "https://www.steelprodukt.ru"),
        ("Дата подготовки", "13 августа 2026 года"),
        ("Статус", "DRAFT — НЕ УТВЕРЖДЕНО"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        set_font(p.add_run(f"{label}: "), size=11, bold=True)
        set_font(p.add_run(value), size=11, color=ORANGE if label == "Статус" else BLACK)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(18)
    set_cellless_rule(rule)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    set_font(p.add_run("Назначение пакета"), size=13, bold=True, color=BLUE)
    p = doc.add_paragraph()
    add_inline(p, "Редактируемые проекты для проверки руководителем, ответственным за организацию обработки персональных данных и российским юристом. Файл не содержит подписей, фиктивных дат и сведений, созданных задним числом.")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(8)
    shade_paragraph(p, LIGHT_ORANGE)
    set_font(p.add_run("  ВАЖНО: техническая подготовка и наличие этого документа не означают юридического утверждения. До отдельного решения internal-раздел остаётся выключенным.  "), size=10.5, bold=True, color=ORANGE)

    doc.add_page_break()
    p = doc.add_paragraph("Состав пакета", style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    for index, source in enumerate(SOURCE_FILES[1:], start=1):
        title = next(
            line[2:].strip()
            for line in (SOURCE_DIR / source).read_text(encoding="utf-8").splitlines()
            if line.startswith("# ")
        )
        p = doc.add_paragraph(style="Number Compact")
        apply_num(p, doc._legal_number_id, 0)
        add_inline(p, title)


def parse_markdown(doc: Document, source: Path, bullet_abstract: int,
                   number_abstract: int) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    buffer: list[str] = []
    active_list: str | None = None
    active_bullet_id: int | None = None
    active_number_id: int | None = None

    def flush() -> None:
        nonlocal buffer
        if not buffer:
            return
        p = doc.add_paragraph()
        add_inline(p, " ".join(part.strip() for part in buffer))
        buffer = []

    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            flush()
            continue
        if line.startswith("# "):
            flush()
            active_list = None
            p = doc.add_paragraph(line[2:].strip(), style="Heading 1")
            p.paragraph_format.space_before = Pt(0)
            continue
        if line.startswith("## "):
            flush()
            active_list = None
            doc.add_paragraph(line[3:].strip(), style="Heading 2")
            continue
        if line.startswith("### "):
            flush()
            active_list = None
            doc.add_paragraph(line[4:].strip(), style="Heading 3")
            continue
        if line.startswith("**") and line.endswith("**") and len(line) > 4:
            flush()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(8)
            shade_paragraph(p, LIGHT_ORANGE)
            set_font(p.add_run("  " + line[2:-2] + "  "), bold=True, color=ORANGE)
            continue
        list_match = re.match(r"^(\s*)-\s+(.*)$", line)
        if list_match:
            flush()
            level = len(list_match.group(1)) // 2
            text = list_match.group(2)
            p = doc.add_paragraph(style="Checklist" if text.startswith("[ ] ") else "Bullet Compact")
            if text.startswith("[ ] "):
                p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
                p.paragraph_format.first_line_indent = Inches(0)
                add_inline(p, "☐ " + text[4:])
            else:
                if active_list != "bullet" or active_bullet_id is None:
                    active_bullet_id = new_num_instance(doc, bullet_abstract)
                active_list = "bullet"
                apply_num(p, active_bullet_id, level)
                add_inline(p, text)
            continue
        number_match = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if number_match:
            flush()
            level = len(number_match.group(1)) // 2
            p = doc.add_paragraph(style="Number Compact")
            if active_list != "number" or active_number_id is None:
                active_number_id = new_num_instance(doc, number_abstract)
            active_list = "number"
            apply_num(p, active_number_id, level)
            add_inline(p, number_match.group(3))
            continue
        if line.startswith("> "):
            flush()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.right_indent = Inches(0.25)
            shade_paragraph(p, "E8EEF5")
            add_inline(p, line[2:])
            continue
        if re.match(r"^`[^`]+`\s*$", line):
            flush()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            shade_paragraph(p, "F3F6F9")
            add_inline(p, line)
            continue
        buffer.append(line)
    flush()


def build() -> Path:
    for filename in SOURCE_FILES:
        if not (SOURCE_DIR / filename).is_file():
            raise SystemExit(f"Missing legal-package source: {filename}")

    doc = Document()
    configure_styles(doc)
    configure_page(doc.sections[0])
    bullet_id, number_id, bullet_abstract, number_abstract = add_numbering(doc)
    doc._legal_number_id = number_id  # task-local helper for the contents page

    properties = doc.core_properties
    properties.title = "Пакет локальных актов по персональным данным"
    properties.subject = "Проекты для юридического утверждения"
    properties.author = "ООО «ЭНЕРГОАЛЬЯНС»"
    properties.keywords = "персональные данные, локальные акты, проект"
    properties.comments = "DRAFT — НЕ УТВЕРЖДЕНО"

    add_cover(doc)
    for index, filename in enumerate(SOURCE_FILES):
        doc.add_page_break()
        p = doc.add_paragraph(style="Document Label")
        p.add_run("ПОЯСНИТЕЛЬНАЯ ЗАПИСКА" if index == 0 else f"ДОКУМЕНТ {index:02d}")
        parse_markdown(doc, SOURCE_DIR / filename, bullet_abstract, number_abstract)

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
